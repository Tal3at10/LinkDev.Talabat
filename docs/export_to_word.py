#!/usr/bin/env python3
"""
سكربت لتحويل ملف Markdown إلى Word
Script to convert Markdown file to Word document
"""

import sys
import os
from pathlib import Path

def convert_markdown_to_word(markdown_file, output_file=None):
    """
    تحويل ملف Markdown إلى Word باستخدام مكتبة python-docx
    
    Args:
        markdown_file (str): مسار ملف Markdown
        output_file (str): مسار ملف Word المخرَج (اختياري)
    """
    try:
        # محاولة استيراد المكتبات المطلوبة
        try:
            import markdown
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as e:
            print(f"خطأ: المكتبات المطلوبة غير مثبتة. {e}")
            print("قم بتثبيت المكتبات باستخدام:")
            print("pip install markdown python-docx")
            return False
        
        # قراءة ملف Markdown
        with open(markdown_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # تحويل Markdown إلى HTML
        html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'codehilite'])
        
        # إنشاء مستند Word جديد
        doc = Document()
        
        # إضافة عنوان المستند
        title = doc.add_heading('شرح مفصل لمشروع LinkDev.Talabat', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # إضافة معلومات المستند
        doc.add_paragraph('تحليل معماري شامل وتوثيق مفصل')
        doc.add_paragraph('تاريخ التحليل: 22 أغسطس 2025')
        doc.add_paragraph('')
        
        # تحويل HTML إلى نص بسيط وإضافته للمستند
        # هذا تحويل بسيط - يمكن تحسينه لمعالجة العناوين والجداول بشكل أفضل
        lines = html_content.split('\n')
        current_heading_level = 0
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # معالجة العناوين
            if line.startswith('<h1>'):
                heading = line.replace('<h1>', '').replace('</h1>', '')
                doc.add_heading(heading, level=1)
                current_heading_level = 1
            elif line.startswith('<h2>'):
                heading = line.replace('<h2>', '').replace('</h2>', '')
                doc.add_heading(heading, level=2)
                current_heading_level = 2
            elif line.startswith('<h3>'):
                heading = line.replace('<h3>', '').replace('</h3>', '')
                doc.add_heading(heading, level=3)
                current_heading_level = 3
            elif line.startswith('<h4>'):
                heading = line.replace('<h4>', '').replace('</h4>', '')
                doc.add_heading(heading, level=4)
                current_heading_level = 4
            elif line.startswith('<h5>'):
                heading = line.replace('<h5>', '').replace('</h5>', '')
                doc.add_heading(heading, level=5)
                current_heading_level = 5
            elif line.startswith('<h6>'):
                heading = line.replace('<h6>', '').replace('</h6>', '')
                doc.add_heading(heading, level=6)
                current_heading_level = 6
            # معالجة الفقرات
            elif line.startswith('<p>'):
                text = line.replace('<p>', '').replace('</p>', '')
                # تنظيف HTML tags البسيطة
                text = text.replace('<strong>', '').replace('</strong>', '')
                text = text.replace('<em>', '').replace('</em>', '')
                text = text.replace('<code>', '').replace('</code>', '')
                doc.add_paragraph(text)
            # معالجة قوائم
            elif line.startswith('<ul>') or line.startswith('</ul>') or line.startswith('<ol>') or line.startswith('</ol>'):
                continue
            elif line.startswith('<li>'):
                text = line.replace('<li>', '').replace('</li>', '')
                p = doc.add_paragraph()
                p.add_run('• ').bold = True
                p.add_run(text)
            # معالجة الكود
            elif line.startswith('<pre>') or line.startswith('</pre>'):
                continue
            elif line.startswith('<code>'):
                text = line.replace('<code>', '').replace('</code>', '')
                p = doc.add_paragraph()
                p.add_run(text).font.name = 'Courier New'
            else:
                # إضافة النص العادي
                if line and not line.startswith('<'):
                    doc.add_paragraph(line)
        
        # حفظ المستند
        if output_file is None:
            output_file = markdown_file.replace('.md', '.docx')
        
        doc.save(output_file)
        print(f"تم تحويل الملف بنجاح إلى: {output_file}")
        return True
        
    except Exception as e:
        print(f"خطأ في التحويل: {e}")
        return False

def main():
    """الدالة الرئيسية"""
    if len(sys.argv) < 2:
        print("الاستخدام: python export_to_word.py <ملف_markdown> [ملف_word_المخرَج]")
        print("مثال: python export_to_word.py Full_Project_Explanation.md")
        return
    
    markdown_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(markdown_file):
        print(f"خطأ: الملف {markdown_file} غير موجود")
        return
    
    success = convert_markdown_to_word(markdown_file, output_file)
    if success:
        print("تم التحويل بنجاح!")
    else:
        print("فشل في التحويل")

if __name__ == "__main__":
    main() 